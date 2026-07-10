"""
Standup Coach Agent

This beginner-friendly script simulates an AI workflow assistant for student
project teams. It reads sample team data and standup responses, identifies
blockers, creates action items, suggests GitHub task board updates, and writes
a formatted Word document summary.
"""

import json
from datetime import date
from pathlib import Path

from docx import Document

from config import (
    get_github_config,
    get_llm_config,
    get_standup_input_mode,
    load_environment,
)
from github_issues import (
    build_github_issue_payloads,
    create_github_issues,
    print_github_issue_dry_run,
    print_github_issue_results,
)
from llm_summary import generate_standup_summary
from slack_parser import slack_messages_to_standup_responses


BASE_DIR = Path(__file__).parent
DATA_DIR = BASE_DIR / "data"
OUTPUT_DIR = BASE_DIR / "output"


def load_json_file(file_path):
    """Load and return JSON data from a file."""
    with open(file_path, "r", encoding="utf-8") as file:
        return json.load(file)


def find_team_member(team, member_id):
    """Find a team member by their ID."""
    for member in team:
        if member["id"] == member_id:
            return member

    return None


def response_has_blocker(response):
    """Return True when a standup response includes a real blocker."""
    blocker_text = response["blockers"].strip().lower()
    no_blocker_answers = ["none", "no", "no blockers", "n/a", "na"]

    return blocker_text not in no_blocker_answers


def identify_blockers(team, responses):
    """Create a list of blockers with the team member attached."""
    blockers = []

    for response in responses:
        if response_has_blocker(response):
            member = find_team_member(team, response["member_id"])
            blockers.append(
                {
                    "name": member["name"],
                    "role": member["role"],
                    "github_username": member["github_username"],
                    "blocker": response["blockers"],
                }
            )

    return blockers


def generate_action_items(blockers):
    """Turn blockers into clear action items for the project team."""
    action_items = []

    for blocker in blockers:
        action_items.append(
            f"Follow up with {blocker['name']} about blocker: {blocker['blocker']}"
        )

    if not action_items:
        action_items.append("No blockers reported today. Keep the current sprint moving.")

    return action_items


def suggest_github_updates(team, responses, blockers):
    """Suggest task board updates without using the real GitHub API."""
    updates = []
    blocked_names = {blocker["name"] for blocker in blockers}

    for response in responses:
        member = find_team_member(team, response["member_id"])

        if member["name"] in blocked_names:
            updates.append(
                {
                    "task_owner": member["github_username"],
                    "suggested_column": "Blocked",
                    "note": f"Review blocker: {response['blockers']}",
                }
            )
        else:
            updates.append(
                {
                    "task_owner": member["github_username"],
                    "suggested_column": "In Progress",
                    "note": f"Continue work: {response['today']}",
                }
            )

    return updates


def load_standup_responses(team, input_mode):
    """Load standup responses from structured JSON or Slack-style messages."""
    if input_mode == "slack":
        slack_messages = load_json_file(DATA_DIR / "sample_slack_messages.json")
        return slack_messages_to_standup_responses(team, slack_messages)

    return load_json_file(DATA_DIR / "sample_standup_responses.json")


def create_docx_summary(
    team, responses, blockers, action_items, github_updates, standup_summary
):
    """Create a formatted Word document summary in the output folder."""
    OUTPUT_DIR.mkdir(exist_ok=True)

    today = date.today().isoformat()
    output_path = OUTPUT_DIR / f"standup_summary_{today}.docx"

    document = Document()
    document.add_heading("Daily Standup Summary", level=1)
    document.add_paragraph(f"Date: {today}")

    document.add_heading("AI-Generated Summary", level=2)
    document.add_paragraph(standup_summary["summary"])
    document.add_paragraph(f"Source: {standup_summary['note']}")

    document.add_heading("Team Updates", level=2)
    for response in responses:
        member = find_team_member(team, response["member_id"])

        document.add_heading(f"{member['name']} - {member['role']}", level=3)
        document.add_paragraph(f"Yesterday: {response['yesterday']}")
        document.add_paragraph(f"Today: {response['today']}")
        document.add_paragraph(f"Blockers: {response['blockers']}")

    document.add_heading("Blockers", level=2)
    if blockers:
        for blocker in blockers:
            document.add_paragraph(
                f"{blocker['name']} ({blocker['role']}): {blocker['blocker']}",
                style="List Bullet",
            )
    else:
        document.add_paragraph("No blockers reported today.")

    document.add_heading("Action Items", level=2)
    for item in action_items:
        document.add_paragraph(item, style="List Number")

    document.add_heading("Suggested GitHub Task Board Updates", level=2)
    for update in github_updates:
        document.add_paragraph(
            f"@{update['task_owner']} -> {update['suggested_column']}: {update['note']}",
            style="List Bullet",
        )

    document.save(output_path)
    return output_path


def main():
    """Run the standup workflow from sample JSON files."""
    load_environment()

    team = load_json_file(DATA_DIR / "sample_team.json")
    input_mode = get_standup_input_mode()
    responses = load_standup_responses(team, input_mode)

    blockers = identify_blockers(team, responses)
    action_items = generate_action_items(blockers)
    github_updates = suggest_github_updates(team, responses, blockers)
    github_config = get_github_config()
    github_issue_payloads = build_github_issue_payloads(blockers, action_items)
    standup_summary = generate_standup_summary(
        team, responses, blockers, action_items, get_llm_config()
    )
    output_path = create_docx_summary(
        team, responses, blockers, action_items, github_updates, standup_summary
    )

    print("Standup Coach Agent finished successfully.")
    print(f"Input mode: {input_mode}")
    print(f"Blockers found: {len(blockers)}")
    print(f"LLM summary used: {standup_summary['used_llm']}")
    print(f"Summary saved to: {output_path}")

    if github_config["create_issues"]:
        github_result = create_github_issues(github_issue_payloads, github_config)
        print_github_issue_results(github_result)
    else:
        print_github_issue_dry_run(github_issue_payloads)


if __name__ == "__main__":
    main()
